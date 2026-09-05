# pyrefly: ignore [missing-import]
"""Resume parsing and ATS match-score helpers.

Pure, rule-based logic (no LLM). PDF/DOCX/TXT text is extracted and normalized,
skills found inside it are compared against the job's required skills/description
to produce a transparent 0-100 match score.
"""
import io
import re
from typing import List

# ---------------------------------------------------------------------------
# Skill extraction / normalization
# ---------------------------------------------------------------------------

# A curated vocabulary of commonly required skills. Matching is case-insensitive
# and tolerant to common punctuation/hyphenation variants ("node-js", "node js").
SKILL_VOCABULARY = [
    # Languages
    "python", "java", "javascript", "typescript", "c++", "c#", "c sharp", "golang",
    "go", "ruby", "php", "swift", "kotlin", "rust", "scala", "dart", "r", "sql",
    "graphql", "shell", "bash", "powershell",
    # Web / Frontend
    "react", "redux", "vue", "angular", "next.js", "nextjs", "svelte", "html",
    "css", "tailwind", "sass", "webpack", "vite", "jquery",
    # Backend / Frameworks
    "node", "node.js", "express", "django", "flask", "fastapi", "spring", "spring boot",
    "rails", "laravel", "asp.net", ".net", "dotnet",
    # Databases & Data
    "postgresql", "postgres", "mysql", "mongodb", "redis", "sqlite", "oracle",
    "elasticsearch", "sql server", "dynamodb", "cassandra", "hadoop", "spark",
    "kafka", "airflow", "pandas", "numpy", "polars", "snowflake", "databricks",
    # ML / AI / Data Science
    "machine learning", "deep learning", "llm", "nlp", "computer vision", "tensorflow",
    "pytorch", "keras", "scikit-learn", "scikit learn", "xgboost", "transformers",
    "langchain", "openai", "hugging face", "mlops", "statistics", "regression",
    "classification", "data analysis", "etl", "data modeling", "data engineering",
    # Cloud / DevOps / Infra
    "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "k8s", "terraform",
    "jenkins", "ci/cd", "github actions", "ansible", "nginx", "linux", "unix", "git",
    "github", "gitlab", "prometheus", "grafana", "serverless", "ec2", "s3", "lambda",
    # Testing / QA
    "pytest", "jest", "mocha", "selenium", "cypress", "junit", "testing",
    "test automation", "unit testing", "tdd",
    # Mobile
    "flutter", "react native", "android", "ios", "kotlin", "swift",
    # Professional / Soft
    "communication", "leadership", "project management", "agile", "scrum", "kanban",
    "teamwork", "problem solving", "analytical thinking", "collaboration",
    # Other
    "rest api", "restful", "microservices", "system design", "oop", "algorithms",
    "data structures", "networking", "cybersecurity", "excel", "power bi", "tableau",
]

# Multi-word phrases must be checked before single words so "machine learning"
# wins over "learning" and "node.js" over "node".
_SKILL_ALIASES = sorted(
    SKILL_VOCABULARY, key=lambda s: len(s), reverse=True
)


def _normalize(text: str) -> str:
    """Lowercase and remove the noise that differs between resumes and job posts."""
    if not text:
        return ""
    text = text.lower()
    # Replace commas/periods/slashes/backslashes with spaces so tokens split
    # cleanly ("react," -> "react"). Keep '+' and '#' so "c++"/"c#" survive, and
    # strip everything else.
    text = re.sub(r"[^a-z0-9+#\s-]", " ", text)
    # Normalize multiple spaces/underscores to a single space.
    text = re.sub(r"[\s_]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def normalize_skill(skill: str) -> str:
    """Return a canonical form of a skill for consistent matching."""
    s = _normalize(skill)
    # Collapse "node js" / "nodejs" / "node" style variants into a stable key is
    # not trivial; instead we keep the vocabulary itself as the canonical source.
    for alias in _SKILL_ALIASES:
        norm = alias.replace("-", " ").replace(".", " ")
        if s == norm or s in (alias, alias.replace("-", " ").replace(".", "")):
            return alias
    return s


def extract_skills_from_text(text: str) -> List[str]:
    """Return the list of known skills present in the given raw text.

    The full vocabulary is scanned over the normalized text so that both resumes
    and job descriptions are parsed against the same canonical set. Word-boundary
    checks are used to avoid false positives (e.g. "go" inside "google").
    """
    norm = _normalize(text)
    if not norm:
        return []
    # Tokenize on spaces; keep only meaningful tokens. This lets us verify each
    # alias as a whole phrase within the normalized resume.
    tokens = norm.split(" ")
    found: List[str] = []

    def token_contains(alias_norm: str) -> bool:
        if not alias_norm:
            return False
        parts = alias_norm.split(" ")
        n = len(parts)
        if n <= 1:
            return parts[0] in tokens
        # Multi-word phrases: look for the contiguous sequence of tokens.
        for i in range(len(tokens) - n + 1):
            if tokens[i : i + n] == parts:
                return True
        return False

    for alias in _SKILL_ALIASES:
        # Reduce hyphens/dots so "node-js", "node js", "nodejs" all collapse.
        reduced = alias.replace("-", " ").replace(".", " ")
        norm_alias = _normalize(reduced)
        if not norm_alias:
            continue
        # Special-handle "c++"/"c#".
        if alias in {"c++", "c#"}:
            if alias == "c++" and any(t in ("c++", "c + +") for t in tokens):
                found.append(alias)
                continue
            if alias == "c#" and "c#" in tokens:
                found.append(alias)
                continue
        # Skip a single-word alias that is covered by a longer alias already
        # matched (e.g. "node" after "node.js").
        if _is_redundant(alias, norm_alias, found):
            continue
        if token_contains(norm_alias):
            if alias not in found:
                found.append(alias)
    return found


def _is_redundant(alias: str, norm_alias: str, found: List[str]) -> bool:
    """True when a shorter alias is subsumed by a longer, already-found alias.

    A single-word alias like "node" is dropped if "node.js" was already matched.
    False-substring matches are avoided by requiring full normalized word tokens.
    """
    norm_alias_words = norm_alias.split(" ")
    # Only dedupe single-word (or single-token) aliases; multi-word phrases are
    # always distinct enough to stand alone.
    if len(norm_alias_words) != 1:
        return False
    word = norm_alias_words[0]
    for matched in found:
        if matched == alias:
            return False
        matched_norm = _normalize(matched.replace("-", " ").replace(".", " "))
        matched_words = matched_norm.split(" ")
        if len(matched_words) < 2:
            continue
        if word in matched_words:
            return True
    return False


# ---------------------------------------------------------------------------
# File text extraction
# ---------------------------------------------------------------------------

MAX_FILE_BYTES = 5 * 1024 * 1024  # 5 MB


class ResumeParseError(Exception):
    """Raised when a resume file cannot be parsed or is invalid."""


def extract_text_from_bytes(data: bytes, extension: str) -> str:
    """Extract plain text from the raw bytes of a resume file.

    Supported: pdf, docx, txt. Legacy binary .doc is rejected with a clear
    message. Raises ResumeParseError for empty, unsupported or corrupt files.
    """
    extension = (extension or "").lower().lstrip(".")
    if len(data) == 0:
        raise ResumeParseError("The uploaded file is empty.")

    if extension == "pdf":
        return _extract_pdf(data)
    if extension == "docx":
        return _extract_docx(data)
    if extension == "txt":
        try:
            return data.decode("utf-8")
        except UnicodeDecodeError:
            return data.decode("latin-1", errors="ignore")
    if extension == "doc":
        # Legacy binary DOC is not reliably parseable here; surface it clearly.
        raise ResumeParseError(
            "Legacy .doc files are not supported. Please convert to .docx or PDF."
        )
    raise ResumeParseError(f"Unsupported file extension: {extension}")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # password-protected, corrupt, etc.
        raise ResumeParseError(f"Could not read PDF: {exc}") from exc
    if reader.is_encrypted:
        raise ResumeParseError("Password-protected PDFs are not supported.")
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            continue
    text = "\n".join(pages).strip()
    if not text:
        raise ResumeParseError("No text could be extracted from the PDF.")
    return text


def _extract_docx(data: bytes) -> str:
    from docx import Document

    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise ResumeParseError(f"Could not read DOCX: {exc}") from exc
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    text = "\n".join(parts).strip()
    if not text:
        raise ResumeParseError("No text could be extracted from the DOCX.")
    return text


# ---------------------------------------------------------------------------
# Match score
# ---------------------------------------------------------------------------

def compute_match_score(resume_text: str, job_skills: list, job_description: str | None = None) -> int:
    """Compute a transparent 0-100 rule-based ATS match score.

    Scoring (fixed weights, fully explainable):
      - Skill matches against the job's required-skill list carry 70% of the weight.
      - Keyword matches against the full job description carry the remaining 30%.

    Returns an integer 0-100; 0 when there is no resume text.
    """
    if not resume_text or not resume_text.strip():
        return 0

    resume_norm = _normalize(resume_text)
    resume_skills = extract_skills_from_text(resume_text)

    skills = list(job_skills or [])
    skills_norm = [_normalize(s) for s in skills]
    skills_norm = [s for s in skills_norm if s]

    description_norm = _normalize(job_description or "")
    description_tokens = [t for t in description_norm.split(" ") if len(t) >= 2]

    # --- Skill component (70%) -------------------------------------------------
    if skills_norm:
        matched_skills = [s for s in skills_norm if _contains(resume_norm, s)]
        skill_ratio = len(matched_skills) / len(skills_norm)
    else:
        # No explicit required skills: reward known skills found in the resume
        # that also appear as keywords in the description, capped at 70%.
        present = [s for s in resume_skills if _contains(resume_norm, s)]
        in_desc = [
            s for s in present
            if _contains(description_norm, s) or any(_contains(s, t) for t in description_tokens)
        ]
        skill_ratio = min(len(in_desc) * 0.12, 1.0)
    skill_component = skill_ratio * 70

    # --- Keyword component (30%) ------------------------------------------------
    if description_tokens:
        matched_tokens = [t for t in description_tokens if _contains(resume_norm, t)]
        keyword_component = (len(matched_tokens) / len(description_tokens)) * 30
    else:
        keyword_component = 0.0

    score = skill_component + keyword_component
    return max(0, min(100, int(round(score))))


def _contains(norm_resume: str, skill: str):
    """True if the (already normalized) skill phrase appears in norm_resume
    as a whole contiguous phrase."""
    f = _normalize(skill)
    if not f:
        return False
    return f" {f} " in f" {norm_resume} "
