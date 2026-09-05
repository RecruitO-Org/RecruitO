"""Unit tests for the rule-based resume parser and ATS match scorer."""
import io
import sys
from pathlib import Path

import pytest

BACKEND_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BACKEND_DIR))

from app.services.resume_parser import (  # noqa: E402
    ResumeParseError,
    compute_match_score,
    extract_skills_from_text,
    extract_text_from_bytes,
    normalize_skill,
)


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def test_extract_empty_file_raises():
    with pytest.raises(ResumeParseError):
        extract_text_from_bytes(b"", "pdf")


def test_extract_unsupported_extension_raises():
    with pytest.raises(ResumeParseError):
        extract_text_from_bytes(b"data", "doc")


def test_extract_txt_utf8():
    text = extract_text_from_bytes("Python Developer\nReact expert".encode("utf-8"), "txt")
    assert "python" in text.lower()
    assert "react" in text.lower()


def test_extract_txt_latin1_fallback():
    # Encoded as latin-1 so utf-8 decode fails, forcing the fallback path.
    data = "Python café".encode("latin-1")
    text = extract_text_from_bytes(data, "txt")
    assert "python" in text.lower()


def test_extract_corrupt_pdf_raises():
    with pytest.raises(ResumeParseError):
        extract_text_from_bytes(b"not a real pdf", "pdf")


@pytest.fixture
def docx_bytes():
    """Build a minimal in-memory DOCX containing known skill keywords."""
    import docx
    buf = io.BytesIO()
    doc = docx.Document()
    doc.add_paragraph("Software Engineer")
    doc.add_paragraph("Proficient in Python and Django")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "aws"
    table.rows[0].cells[1].text = "docker"
    doc.save(buf)
    return buf.getvalue()


def test_extract_docx(docx_bytes):
    text = extract_text_from_bytes(docx_bytes, "docx")
    assert "python" in text.lower()
    assert "django" in text.lower()
    assert "aws" in text.lower()


def test_extract_returns_empty_for_valid_with_no_text():
    # A valid TXT that is only whitespace still returns an empty parse signal.
    text = extract_text_from_bytes("   \n  ".encode(), "txt")
    assert text.strip() == ""


# ---------------------------------------------------------------------------
# Skill normalization & extraction
# ---------------------------------------------------------------------------

def test_normalize_skill():
    assert normalize_skill("React") == "react"
    assert normalize_skill("Node.js") == "node.js"
    assert normalize_skill("  Python  ") == "python"


def test_extract_skills_case_insensitive():
    skills = extract_skills_from_text("PROFICIENT IN PYTHON AND REACT")
    assert "python" in skills
    assert "react" in skills


def test_extract_skills_handles_punctuation():
    skills = extract_skills_from_text("Python, React, and Node.js!")
    assert "python" in skills
    assert "react" in skills
    assert "node.js" in skills


def test_extract_skills_dedup_parent_alias():
    # "node.js" is a superset of "node"; only the longer alias should remain.
    skills = extract_skills_from_text("I use Node.js and Node on the backend")
    assert "node.js" in skills
    assert "node" not in skills


def test_extract_skills_no_false_substring():
    # "c", "go", "r" must not match inside unrelated words.
    skills = extract_skills_from_text("Worked at Google with Cloud and a Ford car")
    assert "go" not in skills
    assert "c" not in skills
    assert "r" not in skills


def test_extract_skills_special_chars():
    skills = extract_skills_from_text("Expert in C++ and C#")
    assert "c++" in skills
    assert "c#" in skills


# ---------------------------------------------------------------------------
# Match score
# ---------------------------------------------------------------------------

SKILLS = ["python", "django", "aws", "kubernetes", "docker"]


def test_match_score_empty_resume_is_zero():
    assert compute_match_score("", SKILLS, "desc") == 0
    assert compute_match_score("   ", SKILLS, "desc") == 0


def test_match_score_no_overlap_is_zero():
    score = compute_match_score("I am fluent in French only", SKILLS, "Needs coding skills")
    assert score == 0


def test_match_score_full_overlap_high():
    resume = "Python Django AWS Kubernetes Docker React"
    score = compute_match_score(resume, SKILLS, "Backend with Python Django AWS Docker")
    assert score >= 70


def test_match_score_partial():
    resume = "5 years Python. React and Node.js experience. AWS cloud and Docker."
    jd = "Python Backend Developer. Requires Python, Django, AWS, Kubernetes, Docker."
    score = compute_match_score(resume, SKILLS, jd)
    assert 20 <= score < 80


def test_match_score_without_skills_list_still_scores():
    # Even with no explicit required-skills, keyword overlap in the description
    # should yield a non-zero score.
    resume = "Python and Django developer with AWS and Docker experience"
    jd = "Experience with Python, Django, AWS and Docker"
    score = compute_match_score(resume, [], jd)
    assert score > 0


def test_match_score_is_bounded_0_100():
    resume = " ".join(SKILLS)
    score = compute_match_score(resume, SKILLS, " ".join(SKILLS))
    assert 0 <= score <= 100
